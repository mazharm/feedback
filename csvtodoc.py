"""
This module converts a CSV file to a Word document
"""
import pandas as pd
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH  # pylint: disable=no-name-in-module


class CSVtoDOCX:
    """
    This class is used convert a CSV file to a Word document
    """

    def __init__(self, summary_dtypes, summary_columns, section_column):
        self.summary_dtypes = summary_dtypes
        self.summary_columns = summary_columns
        self.section_column = section_column
        return

    def convert(self, csv_file, docx_file, heading, workspaces):
        """
        Convert the CSV file to a Word document
        """
        # read the CSV file into a DataFrame
        read_df = pd.read_csv(csv_file, header=0, dtype=self.summary_dtypes)

        # Filter the dataframe to only have rows where the WORKSPACE_OUTPUT column is in the list
        _df = read_df[read_df[self.section_column].isin(workspaces)
                      ][self.summary_columns]

        # Replace NaN values with empty strings in the DataFrame
        _df.fillna('', inplace=True)

        # create a new Word document
        doc = Document()

        # add a heading to the document
        doc.add_heading(heading, 0)

        # loop through each row in the DataFrame
        for idx, row in _df.iterrows():  # pylint: disable=unused-variable

            # add a section header
            section_heading = row.loc[self.section_column]
            doc.add_heading(section_heading, level=1)

            # loop through each column in the row
            for col_name in _df.columns:

                # skip the section column
                if col_name == self.section_column:
                    continue

                # add a subsection header
                subsection_heading = col_name
                doc.add_heading(subsection_heading, level=2)

                # replace '\n' characters with newlines and '[' and ']' with empty strings
                cell_value = str(row[col_name]).replace(
                    '\\n', '\n').replace('[', '').replace(']', '')

                for text in cell_value.split('\n'):
                    # check if the line contains "Key Topics:"
                    if "Key Topics:" in text:
                        # add a bold text paragraph with the line
                        _p = doc.add_paragraph()
                        _p.add_run(text).bold = True
                    else:
                        # add a regular text paragraph with the line
                        _p = doc.add_paragraph(text)
                    # add some formatting to the paragraph
                    _p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    _p.space_after = Inches(0.1)

        # save the Word document
        doc.save(docx_file)
        return


# pylint: disable=pointless-string-statement
"""
This code in place if in case you want to run the tool
as stand alone


summary_dtypes = {'Workspace':str, 'Summary':str, 'Top Quotes':str, 'Action Items':str}

if __name__ == "__main__":
    # create an instance of the CSVtoDOCX class
    csvtodocx = CSVtoDOCX()
    
    # convert the CSV file to a Word document
    csvtodocx.convert(summary_dtypes, 'local/summaries.csv', 'local/summaries.docx', 
                                'NPS Feedback Summary', 'Workspace')
"""
